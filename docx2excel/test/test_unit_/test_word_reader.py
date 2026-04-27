import os
import tempfile
import pytest
from docx import Document

from converters.word_reader import extract_tables_from_word


# =========================
# 工具函数：创建测试用 Word
# =========================
def create_test_docx(tables_data, file_path):
    """
    tables_data:
    [
        [
            ["表头1", "表头2"],
            ["数据1", "数据2"]
        ],
        ...
    ]
    """
    doc = Document()

    for table_data in tables_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))

        for i, row in enumerate(table_data):
            for j, value in enumerate(row):
                table.cell(i, j).text = value

    doc.save(file_path)


# =========================
# 测试1：基础功能
# =========================
def test_extract_single_table_basic():
    tables_data = [
        [
            ["Name", "Age"],
            ["Alice", "18"],
            ["Bob", "20"]
        ]
    ]

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        path = tmp.name

    try:
        create_test_docx(tables_data, path)

        result = extract_tables_from_word(path)

        assert len(result) == 1

        table = result[0]

        # 表头拼接
        assert table["table_title"] == "Name_Age"

        # 数据从第二行开始
        assert table["table_data"] == [
            ["Alice", "18"],
            ["Bob", "20"]
        ]

    finally:
        os.remove(path)


# =========================
# 测试2：去重逻辑（合并单元格场景）
# =========================
def test_extract_table_title_deduplication():
    tables_data = [
        [
            ["标题", "标题", "日期"],
            ["A", "B", "2024"]
        ]
    ]

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        path = tmp.name

    try:
        create_test_docx(tables_data, path)

        result = extract_tables_from_word(path)

        table = result[0]

        # 去重后应该只有一个“标题”
        assert table["table_title"] == "标题_日期"

    finally:
        os.remove(path)


# =========================
# 测试3：多个表格
# =========================
def test_extract_multiple_tables():
    tables_data = [
        [
            ["A", "B"],
            ["1", "2"]
        ],
        [
            ["X", "Y"],
            ["3", "4"]
        ]
    ]

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        path = tmp.name

    try:
        create_test_docx(tables_data, path)

        result = extract_tables_from_word(path)

        assert len(result) == 2

        assert result[0]["table_title"] == "A_B"
        assert result[1]["table_title"] == "X_Y"

    finally:
        os.remove(path)


# =========================
# 测试4：空表（没有数据）
# =========================
def test_extract_empty_table():
    tables_data = [
        []
    ]

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        path = tmp.name

    try:
        # 手动创建一个空文档（无表）
        doc = Document()
        doc.save(path)

        result = extract_tables_from_word(path)

        assert result == []

    finally:
        os.remove(path)


# =========================
# 测试5：空字符串处理
# =========================
def test_extract_with_empty_cells():
    tables_data = [
        [
            ["Name", "Age"],
            ["Alice", ""],
            ["", "20"]
        ]
    ]

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        path = tmp.name

    try:
        create_test_docx(tables_data, path)

        result = extract_tables_from_word(path)

        table = result[0]

        assert table["table_data"] == [
            ["Alice", ""],
            ["", "20"]
        ]

    finally:
        os.remove(path)